import json
from typing import List, Dict, Any
from docx import Document


def normalize_text(text: str) -> str:
    """
    Нормализует текст, убирая лишние пробелы и приводя к читаемому виду.
    """
    return '\n'.join(line.strip() for line in text.replace('\n', '\n').splitlines() if line.strip())


def parse_thematic_plan(file_path: str) -> List[Dict[str, Any]]:
    try:
        # Загрузка DOCX-файла
        doc = Document(file_path)
        tables = doc.tables

        # Ожидаемые заголовки таблицы
        expected_headers = ["неделя",
                            "тема / модуль",
                            "задания", "результат обучения",
                            "литература",
                            "структура оценок"
                            ]

        # Поиск абзаца с заголовком "Тематический план по неделям"
        target_table = None

        for idx, para in enumerate(doc.paragraphs):
            if "тематический план по неделям" in para.text.lower():
                # Ищем таблицу после этого абзаца
                for table in tables:
                    headers = [normalize_text(cell.text) for cell in table.rows[0].cells]
                    if len(headers) == 6 and all(
                            any(exp_header in header.lower() for exp_header in expected_headers) for header in headers):
                        target_table = table
                        break
                break

        if not target_table:
            raise ValueError(
                "Таблица 'Тематический план по неделям' не найдена или не соответствует ожидаемой структуре.")

        # Извлечение заголовков из первой строки таблицы
        headers = [cell.text.strip() for cell in target_table.rows[0].cells]

        # Список для хранения данных
        thematic_plan = []

        # Парсинг строк таблицы (начиная со второй строки)
        for row in target_table.rows[1:]:
            row_data = {}
            for idx, cell in enumerate(row.cells):
                # Нормализуем текст из ячейки
                text = normalize_text('\n'.join([para.text.strip() for para in cell.paragraphs if para.text.strip()]))
                row_data[headers[idx]] = text if text else "Нет данных"

            # Обработка объединенных ячеек (если "Неделя" пустая, берем из предыдущей строки)
            if not row_data.get(headers[0]) and thematic_plan:
                row_data[headers[0]] = thematic_plan[-1][headers[0]]

            thematic_plan.append(row_data)

        return thematic_plan

    except Exception as e:
        print(f"Ошибка: {str(e)}")
        raise


if __name__ == "__main__":
    # Путь к файлу
    file_path = "media/syllabus_files/Силлабус_Философия_2024-2025.docx"
    try:
        # Парсинг таблицы
        thematic_plan_data = parse_thematic_plan(file_path)

        # Преобразование в JSON
        json_data = json.dumps(thematic_plan_data, ensure_ascii=False, indent=4)

        # Вывод результата
        print(json_data)

        # Сохранение в файл
        with open("thematic_plan.json", "w", encoding="utf-8") as f:
            f.write(json_data)
    except Exception as e:
        print(f"Ошибка: {str(e)}")